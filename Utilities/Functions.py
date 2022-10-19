import os
import time
import docx
import openpyxl
from selenium.webdriver.common import keys
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches
from selenium.webdriver import Keys, ActionChains
from Utilities.Config import screenshot_folder, root_dir
from Utilities.Utils import Utilities
from Utilities.WebMisc import WebMisc

log = Utilities().getlogger()

class Functions:

    def bookmark(self, module, ts_id, tc_id, step_num):
        path = screenshot_folder + f'{module}\\{ts_id}'
        document = Document(path + f'\\{ts_id}.docx')
        tc_desc = self.get_tc(module, tc_id, step_num)[0]
        tc_expected = self.get_tc(module, tc_id, step_num)[1]

        if tc_id == 'Precondition':
            table = document.add_table(rows=0, cols=1, style=document.styles['Table Grid'])
            data_row = table.add_row().cells
            data_row[0].text = self.get_precondition(module, ts_id)  # Get_TC_bookmark
            document.add_paragraph()

        else:
            if step_num == 'Step 1':
                table = document.add_table(rows=0, cols=1, style=document.styles['Table Grid'])
                data_row = table.add_row().cells
                data_row[0].text = self.get_tc_name(module, tc_id)  # Get_TC_bookmark
                document.add_paragraph()

            if tc_desc is not None and len(tc_desc) > 0:
                table = document.add_table(rows=0, cols=1, style=document.styles['Table Grid'])

                if step_num != 'Step 1':
                    document.add_page_break()

                data_row = table.add_row().cells
                data_row[0].text = step_num + " - " + tc_desc  # Get_Step_bookmark
                data_row = table.add_row().cells
                data_row[0].text = tc_expected
                log.info(tc_desc)

        document.save(path + f'\\{ts_id}.docx')  # update document

    def screen_capture(self, driver, module, ts_id, tc_id, step_num):
        time.sleep(2)
        path = screenshot_folder + f'{module}\\{ts_id}'
        document = Document(path + f'\\{ts_id}.docx')
        driver.save_screenshot(path + f'/{ts_id}_{tc_id} {step_num}.png')
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(path + f'/{ts_id}_{tc_id} {step_num}.png', width=Inches(5.8))

        document.save(path + f'\\{ts_id}.docx')  # update document

    def tag_status(self, module, ts_id, status):
        path = screenshot_folder + f'{module}\\{ts_id}'
        document = Document(path + f'\\{ts_id}.docx')
        document.add_page_break()
        document.add_paragraph(status)
        document.save(path + f'\\{ts_id}.docx')

    def create_document(self, driver, module, ts_id):
        path = screenshot_folder + f'{module}\\{ts_id}'
        document = docx.Document()

        is_exist = os.path.exists(path)
        if not is_exist:
            os.mkdir(path)

        table = document.add_table(rows=1, cols=1, style=document.styles['Table Grid'])
        row = table.rows[0].cells
        row[0].text = ts_id + " - " + self.get_ts(module, ts_id)
        document.add_paragraph()
        document.save(path + f'\\{ts_id}.docx')

    def get_ts(self, module, ts_id):
        workbook = openpyxl.load_workbook(root_dir.parent/f'Config/{module}.xlsx')
        sheet = workbook["AT - Test Suite"]
        total_rows = sheet.max_row
        data = ""
        for i in range(2, total_rows + 1):
            if sheet.cell(row=i, column=1).value == ts_id:
                data = sheet.cell(row=i, column=2).value
        return data

    def get_tc(self, module, tc_id, step_num):
        workbook = openpyxl.load_workbook(root_dir.parent/f'Config/{module}.xlsx')
        sheet = workbook["AT - Test Case"]
        total_rows = sheet.max_row
        description = ""
        expected = ""

        for i in range(2, total_rows + 1):
            if sheet.cell(row=i, column=1).value == tc_id and sheet.cell(row=i, column=2).value == step_num:
                description = sheet.cell(row=i, column=3).value
                expected = sheet.cell(row=i, column=4).value

        return description, expected

    def get_tc_name(self, module, tc_id):
        workbook = openpyxl.load_workbook(root_dir.parent/f'Config/{module}.xlsx')
        sheet = workbook["AT - Test Case"]
        total_rows = sheet.max_row
        tc_name = ""

        for i in range(2, total_rows + 1):
            if sheet.cell(row=i, column=1).value is not None and tc_id + "_" in sheet.cell(row=i, column=1).value:
                tc_name = sheet.cell(row=i, column=1).value

        return tc_name

    def get_precondition(self, module, ts_id):
        workbook = openpyxl.load_workbook(root_dir.parent/f'Config/{module}.xlsx')
        sheet = workbook["AT - Test Suite"]
        total_rows = sheet.max_row
        data = ""
        column_tag = ts_id + " Precondition"
        for i in range(2, total_rows + 1):
            if sheet.cell(row=i, column=1).value == column_tag:
                data = sheet.cell(row=i, column=2).value
        return data

    def new_tab(self, driver, url):
        driver.execute_script("window.open('');")
        driver.switch_to.window(driver.window_handles[1])
        driver.get(url)
        log.info("Opened: " + url)
        curr = driver.current_window_handle
        for handle in driver.window_handles:
            driver.switch_to.window(handle)
            if handle != curr:
                driver.close()

    def input_text(self, element, value):
        element.send_keys(value)
        log.info("input: "+element.text + value)

    def enter(self, element):
        element.send_keys(Keys.ENTER)
        log.info("Enter: "+element.text)

    def page_down(self, driver, count):
        for x in range(count):
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.PAGE_DOWN)
            log.info("Page Down")
            time.sleep(1)

    def coordinates_click(self, driver, x, y):
        body = driver.find_element(By.TAG_NAME, "body")
        action = ActionChains(driver).move_to_element_with_offset(body, x, y)
        action.click()
        action.perform()

    def scroll_element(self, driver, element):
        actions = ActionChains(driver)
        actions.move_to_element(element).perform()

    def full_enter(self, driver):
        driver.Keyboard.PressKey(Keys.ENTER)
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.ENTER)
        log.info("Entered")

    def click(self, element):
        WebMisc().is_enabled(element)
        element.click()
        log.info("Clicked")
        #try:
        #    log.info(element.text + " clicked")
        #except:
        #    log.info("Clicked")

    def option_click(self, element):
        if element is not None:
            WebMisc().is_enabled(element)
            element.click()
            try:
                log.info(element.text + " clicked")
            except:
                log.info("Clicked")

    def accept_alert(self, driver):
        time.sleep(4)
        driver.switch_to.alert.accept()

    def modal_click(self, driver, element):
        if element is not None:
            WebMisc().is_enabled(element)
            driver.execute_script("arguments[0].click();", element)
            try:
                log.info(element.text + " clicked")
            except:
                log.info("Clicked")

    def verify(self, element):
        if element.is_displayed():
            log.info(element.text + " exists")

    def verify_text(self, element, text):
        assert element.text == text
        log.info(element.text + " exists")
